import os
import sys
import json
from docx import Document

# Add backend to path
sys.path.append(os.path.join(os.getcwd(), 'backend'))

from resume_generator import get_optimization_plan, generate_tailored_resume, _get_paragraph_info

def verify():
    # 1. Setup paths
    resume_path = "backend/uploads/VEEKSHITH GULLAPUDI_RESUME_main.docx"
    output_path = "backend/generated/test_verified_resume.docx"
    job_desc = "ServiceNow Developer at Vlink.inc. Requirements: JavaScript, ITIL, ServiceNow ITSM, leadership."
    
    if not os.path.exists(resume_path):
        print(f"ERROR: Resume not found at {resume_path}")
        return

    print(f"--- 1. Testing Paragraph Detection ---")
    doc = Document(resume_path)
    experience_headers = ["experience", "professional experience"]
    summary_headers = ["summary", "profile"]
    stop_headers = ["education", "skills"]
    
    for i, p in enumerate(doc.paragraphs[:20]):
        p_type, text = _get_paragraph_info(p, experience_headers, summary_headers, stop_headers)
        if text:
            print(f"[{p_type}] {text[:50]}...")

    print(f"\n--- 2. Simulating Optimization Plan (Mocked AI) ---")
    # We call the real logic but we'll mock the AI call inside if possible, 
    # but since it's hard to mock without modifying source, 
    # we'll just manually construct a plan that simulates the problem cases.
    
    mock_plan = {
        "company_name": "Vlink_inc",
        "summary": {
            "original": "Results-oriented Professional...",
            "optimized": "Highly skilled ServiceNow Developer with expertise in JavaScript and ITIL..."
        },
        "experience_entries": [
            {
                "header": "Full Stack Java Developer | Morgan Stanley",
                "bullets": ["Existing bullet 1", "Existing bullet 2"],
                "optimized_bullets": [
                    "• Spearheaded ServiceNow implementation for Morgan Stanley...",
                    "• Optimized ITIL workflows reducing ticket resolution time by 20%.",
                    "• NEW: Architected custom ServiceNow portals for cross-functional teams."
                ]
            },
            {
                "header": "Project Engineer | Wipro",
                "bullets": ["Old Wipro bullet"],
                "optimized_bullets": [
                    "• Delivered high-quality Java solutions at Wipro.",
                    "• NEW: Integrated ServiceNow APIs with legacy Java systems."
                ]
            }
        ]
    }

    print(f"--- 3. Testing Final Generation with Mock Plan ---")
    final_path, msg, company = generate_tailored_resume(
        resume_path, 
        job_desc, 
        output_path, 
        api_key=None, 
        approved_plan=mock_plan
    )
    
    print(f"Message: {msg}")
    print(f"Detected Company: {company}")
    print(f"Output saved to: {final_path}")

    if os.path.exists(final_path):
        print(f"\n--- 4. Verifying Output Structure ---")
        out_doc = Document(final_path)
        bullet_count = 0
        for p in out_doc.paragraphs:
            # Check for double bullets and content
            if "• •" in p.text:
                print(f"!! FAILURE: Double bullet found: {p.text}")
            if "Vlink_inc" in p.text:
                print(f"Detected correct company injection.")
            
            # Simple check for our mock bullets
            if "ServiceNow implementation" in p.text:
                print(f"Success: Correctly replaced Morgan Stanley bullet.")
                bullet_count += 1
            if "Wipro" in p.text:
                bullet_count += 1
        
        print(f"Verified {bullet_count} key changes in the output document.")
        
        # Check for trailing empty bullets (simple check: any paragraph with JUST a bullet char)
        empty_bullets = 0
        for p in out_doc.paragraphs:
            if p.text.strip() in ["•", "-", "*", "·"]:
                empty_bullets += 1
        
        if empty_bullets > 0:
            print(f"!! FAILURE: Found {empty_bullets} empty bullet points.")
        else:
            print("Success: No empty trailing bullets found.")

if __name__ == "__main__":
    verify()
