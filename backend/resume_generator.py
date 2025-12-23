# Copyright (c) 2025 Veekshith Gullapudi. All rights reserved.

from docx import Document
import os
import re

# Support both Lambda (relative) and local dev (absolute) imports
try:
    from .copilot import ResumeCopilot
except ImportError:
    from copilot import ResumeCopilot

def _get_paragraph_info(paragraph, experience_headers, summary_headers, stop_headers):
    """ Helper to categorize a paragraph. """
    text = paragraph.text.strip()
    if not text:
        return "empty", ""
    
    lower_text = text.lower()
    is_bold = paragraph.runs and paragraph.runs[0].bold
    is_heading_style = paragraph.style.name.startswith('Heading')
    is_caps = len(text) < 50 and text.isupper()
    
    # 1. Section Headers
    if is_heading_style or is_caps or (is_bold and len(text) < 50):
        if any(h in lower_text for h in experience_headers):
            return "experience_section_header", text
        if any(h in lower_text for h in summary_headers):
            return "summary_section_header", text
        if any(sh in lower_text for sh in stop_headers):
            return "stop_section_header", text
            
    # 2. Bullets
    is_bullet = False
    try:
        # Check XML numbering, style names, or common bullet characters
        if (paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None) or \
           'list' in paragraph.style.name.lower() or text.startswith(('•', '-', '–', '*', '·', '◦')):
            is_bullet = True
    except: pass
    
    if is_bullet:
        return "bullet", text
        
    # 3. Entry Headers (everything else inside Experience that isn't a bullet/section header)
    return "text", text

def get_optimization_plan(original_path: str, job_description: str, api_key: str = None) -> dict:
    """
    Analyzes the resume and returns a plan of suggested changes (Summary + Grouped Experience).
    Used for the 'Preview' feature.
    """
    print(f"DEBUG: Generating optimization plan for {original_path}")
    doc = Document(original_path)
    copilot = ResumeCopilot(api_key)
    
    experience_headers = ["experience", "professional experience", "work history", "employment history", "work experience", "career history"]
    summary_headers = ["summary", "professional summary", "objective", "profile", "professional profile"]
    stop_headers = ["education", "skills", "projects", "technical skills", "certifications", "awards", "languages", "volunteering", "top skills"]
    
    plan = {
        "company_name": "Company",
        "summary": {"original": "", "optimized": ""},
        "experience_entries": []
    }
    
    current_section = None # None, "summary", "experience"
    current_entry = None
    
    for para in doc.paragraphs:
        p_type, text = _get_paragraph_info(para, experience_headers, summary_headers, stop_headers)
        
        if p_type == "empty": continue
        
        if p_type == "summary_section_header":
            current_section = "summary"
        elif p_type == "experience_section_header":
            current_section = "experience"
        elif p_type == "stop_section_header":
            current_section = "stop"
        elif current_section == "summary" and p_type == "text":
            plan["summary"]["original"] += text + " "
        elif current_section == "experience":
            if p_type == "bullet":
                if not current_entry:
                    # Rare: bullets without a header? Create a placeholder entry
                    current_entry = {"header": "Experience", "bullets": []}
                    plan["experience_entries"].append(current_entry)
                current_entry["bullets"].append(text)
            elif p_type == "text":
                # Entry Header
                # If we were already in an entry and it has bullets, start a new one.
                # If it doesn't have bullets, append to current header (multi-line header).
                if not current_entry or current_entry["bullets"]:
                    current_entry = {"header": text, "bullets": []}
                    plan["experience_entries"].append(current_entry)
                else:
                    current_entry["header"] += " | " + text

    # 2. Call AI in bulk to avoid rate limits
    batch_result = copilot.optimize_all_content(plan["summary"]["original"], plan["experience_entries"], job_description)
    
    if batch_result:
        # Map batch result back to plan
        if "company_name" in batch_result:
            plan["company_name"] = batch_result["company_name"]

        if "summary" in batch_result:
            plan["summary"]["optimized"] = batch_result["summary"]
        
        if "entries" in batch_result:
            for entry_res in batch_result["entries"]:
                idx = entry_res.get("id")
                if idx is not None and idx < len(plan["experience_entries"]):
                    plan["experience_entries"][idx]["optimized_bullets"] = entry_res.get("optimized_bullets", [])
    else:
        # Fallback if batch fails
        print("DEBUG: Batch optimization failed, fallback to originals.")
        plan["summary"]["optimized"] = plan["summary"]["original"]
        for entry in plan["experience_entries"]:
            entry["optimized_bullets"] = entry["bullets"]

    return plan

def generate_tailored_resume(original_path: str, job_description: str, output_path: str, api_key: str = None, approved_plan: dict = None) -> tuple[str, str, str]:
    """
    Generates the final file. If approved_plan is provided, uses that.
    Otherwise, generates a plan on the fly.
    """
    try:
        doc = Document(original_path)
        copilot = ResumeCopilot(api_key)
        
        plan = approved_plan or get_optimization_plan(original_path, job_description, api_key)
        
        # Priority: approved_plan company name, then from analysis
        company_name = plan.get("company_name", "Company")
        
        # Mapping logic back to doc is complex because we want to REPLACE summary and APPEND/REPLACE bullets
        # Since this is a specialized task, I'll implement a clean-replacement strategy:
        # 1. Find the summary and experience sections again.
        # 2. Update their paragraphs.
        
        experience_headers = ["experience", "professional experience", "work history", "employment history", "work experience", "career history"]
        summary_headers = ["summary", "professional summary", "objective", "profile", "professional profile"]
        stop_headers = ["education", "skills", "projects", "technical skills", "certifications", "awards", "languages", "volunteering", "top skills"]

        current_section = None
        exp_entry_idx = -1
        bullet_idx = 0
        last_bullet_para = None
        paragraphs_to_delete = []
        
        summary_paras = []
        
        for para in doc.paragraphs:
            p_type, text = _get_paragraph_info(para, experience_headers, summary_headers, stop_headers)
            
            if p_type == "empty": continue
            
            if p_type == "summary_section_header":
                current_section = "summary"
                continue
            elif p_type == "experience_section_header":
                current_section = "experience"
                exp_entry_idx = -1
                continue
            elif p_type == "stop_section_header":
                # Finalize last entry before leaving
                if current_section == "experience" and exp_entry_idx >= 0 and last_bullet_para:
                    _append_extra_bullets(plan["experience_entries"][exp_entry_idx], bullet_idx, last_bullet_para)
                current_section = "stop"
                continue

            if current_section == "summary" and p_type == "text":
                summary_paras.append(para)
                
            elif current_section == "experience":
                if p_type == "text":
                    # Transition to next entry if the current one had bullets
                    # THIS MUST MATCH THE LOGIC IN get_optimization_plan
                    # Entries are created when we hit 'text' if (first entry OR previous entry has bullets)
                    was_in_entry = exp_entry_idx >= 0
                    has_bullets_in_prev = was_in_entry and bullet_idx > 0
                    
                    if not was_in_entry or has_bullets_in_prev:
                        # Append extras to previous before starting new
                        if was_in_entry and last_bullet_para:
                            _append_extra_bullets(plan["experience_entries"][exp_entry_idx], bullet_idx, last_bullet_para)
                        
                        exp_entry_idx += 1
                        bullet_idx = 0
                        last_bullet_para = None
                    
                elif p_type == "bullet":
                    if exp_entry_idx == -1: exp_entry_idx = 0
                    
                    if exp_entry_idx < len(plan["experience_entries"]):
                        entry = plan["experience_entries"][exp_entry_idx]
                        if bullet_idx < len(entry["optimized_bullets"]):
                            # Detect bullet style
                            is_word_bullet = False
                            try:
                                if (para._p.pPr is not None and para._p.pPr.numPr is not None) or 'list' in para.style.name.lower():
                                    is_word_bullet = True
                            except: pass

                            # Find existing literal bullet character
                            match = re.match(r'^(\s*[•\-\–*·◦]\s*)', para.text)
                            
                            new_text = entry["optimized_bullets"][bullet_idx]
                            clean_new_text = re.sub(r'^(\s*[•\-\–*·◦]\s*)', '', new_text.strip())

                            if match:
                                # Preserve the literal bullet the user had
                                para.text = match.group(1) + clean_new_text
                            elif is_word_bullet:
                                # Word is handling the bullet, don't add characters
                                para.text = clean_new_text
                            else:
                                # Fallback: add a standard bullet
                                para.text = "• " + clean_new_text
                            
                            bullet_idx += 1
                            last_bullet_para = para
                        else:
                            # Mark for deletion if AI shortened the list
                            paragraphs_to_delete.append(para)

        # Final cleanup for the very last entry
        if current_section == "experience" and exp_entry_idx >= 0 and exp_entry_idx < len(plan["experience_entries"]) and last_bullet_para:
            _append_extra_bullets(plan["experience_entries"][exp_entry_idx], bullet_idx, last_bullet_para)

        # Apply summary changes
        if summary_paras and plan["summary"]["optimized"]:
            summary_paras[0].text = plan["summary"]["optimized"]
            for p in summary_paras[1:]:
                paragraphs_to_delete.append(p)

        # Cleanup: Remove paragraphs marked for deletion
        for p in paragraphs_to_delete:
            try:
                element = p._element
                element.getparent().remove(element)
            except: pass

        doc.save(output_path)
        if not company_name or company_name == "Company":
            company_name = copilot.extract_company_name(job_description)
        
        return output_path, "Resume generated successfully", company_name

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        return "", f"Error: {e}", ""

def _append_extra_bullets(entry, current_idx, anchor_para):
    """ Helper to append extra optimized bullets if AI generated more than were original. """
    if "optimized_bullets" in entry and current_idx < len(entry["optimized_bullets"]):
        for j in range(current_idx, len(entry["optimized_bullets"])):
            new_bullet = entry["optimized_bullets"][j]
            # Strip existing bullet char from new_bullet to avoid double bullets
            clean_bullet = re.sub(r'^(\s*[•\-\–*·◦]\s*)', '', new_bullet.strip())
            
            # Check if anchor preferred a literal bullet or has a list style
            match = re.match(r'^(\s*[•\-\–*·◦]\s*)', anchor_para.text)
            is_word_bullet = False
            try:
                if (anchor_para._p.pPr is not None and anchor_para._p.pPr.numPr is not None) or 'list' in anchor_para.style.name.lower():
                    is_word_bullet = True
            except: pass

            prefix = "\n• "
            if match:
                prefix = "\n" + match.group(1)
            elif is_word_bullet:
                prefix = "\n"
            
            anchor_para.text += prefix + clean_bullet
